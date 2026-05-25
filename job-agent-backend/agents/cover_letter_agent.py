"""
Cover Letter Agent - Generates professional cover letters using Groq LLM.
"""
import json
import logging
import os
from pathlib import Path
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, Any
from utils.file_helpers import sanitise_filename


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Groq client
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY)


def generate_cover_letter(job: Dict[str, Any], summary: str, output_dir: str) -> str:
    """
    Generate a professional cover letter for a job posting.
    
    Args:
        job: Job dictionary with keys: title, company, location, description, job_url
        summary: Candidate's professional summary/background
        output_dir: Directory to save the cover letter
    
    Returns:
        Full file path to the generated cover letter (.docx file)
    """
    try:
        # Call Groq API to generate cover letter
        prompt = f"""Write a concise, professional cover letter (under 250 words, 3 paragraphs) for a job application.

Job Information:
- Title: {job.get('title', '')}
- Company: {job.get('company', '')}
- Description: {job.get('description', '')}

Candidate Summary:
{summary}

Requirements:
- Keep it professional and compelling
- 3 paragraphs only
- Under 250 words total
- Start with an introduction mentioning the position and company
- Middle paragraph highlighting relevant qualifications
- Closing paragraph expressing enthusiasm

Return ONLY the cover letter text with no additional formatting or explanation."""
        
        message = client.messages.create(
            model="llama3-8b-8192",
            max_tokens=1500,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )
        
        cover_letter_text = message.content[0].text.strip()
        
        # Create .docx document
        doc = Document()
        
        # Add title
        company = job.get("company", "Company")
        title = job.get("title", "Position")
        doc.add_heading(f"Cover Letter for {title} at {company}", 0)
        
        # Add cover letter content
        for paragraph_text in cover_letter_text.split("\n\n"):
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text.strip())
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Sanitise filename
        sanitised_company = sanitise_filename(company)
        sanitised_title = sanitise_filename(title)
        filename = f"{sanitised_company}_{sanitised_title}_cover_letter.docx"
        
        # Ensure output directory exists
        output_path = Path(output_dir) / filename
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Save document
        doc.save(str(output_path))
        
        logger.info(f"Generated cover letter: {output_path}")
        return str(output_path)
    
    except Exception as e:
        logger.error(f"Error generating cover letter: {e}")
        return ""
