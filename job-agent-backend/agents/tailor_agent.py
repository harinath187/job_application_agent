"""
Tailor Agent - Tailors resume content to match job requirements using Groq LLM.
"""
import json
import logging
import os
from pathlib import Path
from groq import Groq
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, List, Any
from utils.file_helpers import sanitise_filename


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Groq client
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
client = Groq(api_key=GROQ_API_KEY)


def tailor_resume(resume_text: str, job: Dict[str, Any], skills: List[str]) -> Dict[str, Any]:
    """
    Tailor resume content to match a specific job posting.
    
    Args:
        resume_text: Full text from the candidate's resume
        job: Job dictionary with keys: title, company, location, description, job_url
        skills: List of extracted skills from resume
    
    Returns:
        Dictionary with keys: rewritten_summary (str), revised_skills (List[str]), 
        bullet_rewrites (List[str])
    """
    try:
        # Truncate job description to 800 characters for prompt efficiency
        job_description_snippet = job.get('description', '')[:800]
        
        # Call Groq API to tailor resume
        prompt = f"""You are a professional resume writer. Tailor the following resume to match the job posting.

IMPORTANT: Do NOT fabricate experience or skills. Only rewrite existing content to emphasize relevant qualifications.

Job Title: {job.get('title', '')}
Company: {job.get('company', '')}
Job Description:
{job_description_snippet}

Current Resume:
{resume_text}

Current Extracted Skills: {', '.join(skills)}

Provide a JSON response with:
1. rewritten_summary: A 2-3 sentence professional summary tailored to this job (emphasizing relevant background)
2. revised_skills: A list of relevant skills from the resume and provided skills, prioritized for this role (5-8 items). Only include skills that appear in either the Current Resume text or the Current Extracted Skills list. Do not add skills not present in either source.
3. bullet_rewrites: Identify the 3 experience bullets from the resume most relevant to this job, then rewrite each one to better match the job description's language — without changing the underlying achievement or adding new experience.

IMPORTANT: Mirror the exact keywords and phrases from the job description where they honestly reflect the candidate's experience. This improves ATS compatibility.

If there are fewer than 3 experience bullets in the resume, return however many exist. Never return null for any field — use an empty list [] if needed.

Return ONLY valid JSON with no additional text. Example format:
{{"rewritten_summary": "...", "revised_skills": [...], "bullet_rewrites": [...]}}"""
        
        message = client.messages.create(
            model="llama3-8b-8192",
            max_tokens=1500,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )
        
        response_text = message.content[0].text.strip()
        
        # Parse JSON response
        tailored_data = json.loads(response_text)
        
        return {
            "rewritten_summary": str(tailored_data.get("rewritten_summary", "")),
            "revised_skills": tailored_data.get("revised_skills", skills) if isinstance(tailored_data.get("revised_skills"), list) else skills,
            "bullet_rewrites": tailored_data.get("bullet_rewrites", []) if isinstance(tailored_data.get("bullet_rewrites"), list) else []
        }
    
    except json.JSONDecodeError as e:
        logger.error(f"Failed to parse JSON response from Groq: {e}")
        return {
            "rewritten_summary": "",
            "revised_skills": skills,
            "bullet_rewrites": []
        }
    except Exception as e:
        logger.error(f"Error tailoring resume: {e}")
        return {
            "rewritten_summary": "",
            "revised_skills": skills,
            "bullet_rewrites": []
        }


def save_tailored_resume(resume_text: str, tailored_data: Dict[str, Any], job: Dict[str, Any], output_dir: str) -> str:
    """
    Save a tailored resume as a DOCX file.
    
    Args:
        resume_text: Original resume text
        tailored_data: Tailored content with rewritten_summary, revised_skills, bullet_rewrites
        job: Job dictionary with title and company
        output_dir: Directory to save the tailored resume
    
    Returns:
        Full file path to the saved tailored resume
    """
    try:
        # Create .docx document
        doc = Document()
        
        # Add title
        company = job.get("company", "Company")
        title = job.get("title", "Position")
        doc.add_heading(f"Resume - {title} at {company}", 0)
        
        # Add professional summary
        if tailored_data.get("rewritten_summary"):
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(tailored_data.get("rewritten_summary"))
        
        # Add skills section
        if tailored_data.get("revised_skills"):
            doc.add_heading("Key Skills", level=1)
            skills_text = ", ".join(tailored_data.get("revised_skills", []))
            doc.add_paragraph(skills_text)
        
        # Add experience highlights
        if tailored_data.get("bullet_rewrites"):
            doc.add_heading("Experience Highlights", level=1)
            for bullet in tailored_data.get("bullet_rewrites", []):
                p = doc.add_paragraph(bullet, style="List Bullet")
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Add original resume content for reference
        doc.add_heading("Full Resume", level=1)
        for paragraph_text in resume_text.split("\n"):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text.strip())
        
        # Sanitise filename
        sanitised_company = sanitise_filename(company)
        sanitised_title = sanitise_filename(title)
        filename = f"{sanitised_company}_{sanitised_title}_tailored_resume.docx"
        
        # Ensure output directory exists
        output_path = Path(output_dir) / filename
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Save document
        doc.save(str(output_path))
        
        logger.info(f"Generated tailored resume: {output_path}")
        return str(output_path)
    
    except Exception as e:
        logger.error(f"Error saving tailored resume: {e}")
        return ""
