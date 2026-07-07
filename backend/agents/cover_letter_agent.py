"""
Cover Letter Agent - Generates professional, tailored cover letters using Groq LLM.
Features structured JD parsing, keyword extraction, and fallback templates.
"""
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, Any
from utils.groq_client import groq_call
from utils.file_helpers import sanitise_filename


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def build_advanced_prompt(
    company_name: str,
    job_title: str,
    job_description: str,
    tailored_resume_summary: str,
    extracted_skills: list
) -> str:
    """
    Build an advanced cover letter prompt with structured JD parsing and keyword extraction.
    
    Args:
        company_name: Target company name
        job_title: Position title from JD
        job_description: Full job description text
        tailored_resume_summary: Candidate's tailored resume summary (from Resume Tailor Agent)
        extracted_skills: List of candidate skills extracted from resume PDF
    
    Returns:
        Optimized prompt string ready for Groq LLM
    """
    skills_str = ", ".join(extracted_skills) if extracted_skills else "relevant technical skills"
    resume_summary_str = tailored_resume_summary or "strong technical background"
    
    return f"""You are an expert career coach writing a highly tailored, specific cover letter.

=== CANDIDATE PROFILE ===
Resume Summary: {resume_summary_str}
Core Skills: {skills_str}

=== JOB POSTING ===
Position: {job_title}
Company: {company_name}
Description:
{job_description}

=== YOUR TASK ===
Write a concise 3-paragraph cover letter (150–220 words total) that is specific to this job.

**STEP 1: Parse the job description and extract:**
- Required or preferred certifications/credentials
- Specific tools mentioned (e.g., Axe, WAVE, Jira, Figma, Salesforce)
- Standards, frameworks, or methodologies (e.g., WCAG, Agile, CI/CD)
- 3–4 core job responsibilities or challenges
- Technical skills explicitly listed

**STEP 2: Map candidate skills to JD requirements:**
- Identify 2–3 of the candidate's skills that directly match or exceed JD needs
- Note any certifications or tools the candidate has that the JD values
- Identify 1 skill the candidate has that differentiates them

**STEP 3: Write the cover letter with EXACT paragraph structure:**

Paragraph 1 (Opening – ~50 words):
- State the specific position and company name
- Reference ONE specific JD requirement, tool, or responsibility that excites you
- Example: "I'm excited to apply for the Accessibility Engineer role at [Company], where WCAG 2.1 compliance and JAWS testing are core to ensuring inclusive experiences."
- NO generic openings like "I am excited to apply" by itself

Paragraph 2 (Skills Match – ~100 words):
- Pick 2–3 of the candidate's skills that appear in the JD
- For EACH skill, give a SPECIFIC example from the tailored resume summary
- Name specific tools from the JD that the candidate has used
- Reference certifications they hold if relevant to the JD
- Avoid "I am a team player" or "I have excellent communication skills"
- Example: "My 3 years testing with Axe and WAVE have given me deep expertise with WCAG 2.1 compliance. I've led accessibility audits on 15+ projects using NVDA and JAWS—the tools [Company] relies on."

Paragraph 3 (Closing – ~50 words):
- Restate one differentiator or tool match from Paragraph 2
- Make a clear, specific call to action
- Reference the company by name
- NO hollow phrases like "I look forward to hearing from you"
- Example: "I'm eager to bring my Lighthouse expertise and Section 508 knowledge to [Company]'s mission. I'd welcome a conversation about how I can contribute."

=== OUTPUT FORMAT ===
Output ONLY the 3-paragraph letter body. No subject line, no salutation, no signature, no metadata.
Separate paragraphs with a single blank line.
No markdown formatting — plain text only.

=== STRICT CONSTRAINTS ===

**DO:**
✓ Use words, tools, and certifications directly from the JD
✓ Reference the candidate's specific resume summary points
✓ Tie each paragraph to a concrete JD requirement
✓ Use active voice and specific examples
✓ Keep sentences under 20 words when possible
✓ Paraphrase the resume summary; do not copy verbatim

**DO NOT:**
✗ Use generic fillers: "relevant skills", "excited to", "passionate about", "dynamic", "synergy", "team player"
✗ Include subject lines, salutations, signatures, or metadata
✗ Paste raw job description text into the letter
✗ Include unfilled placeholders like {{name}} or [TOOL]
✗ Repeat the same word or fact twice
✗ Apologize or hedge ("I hope", "I think", "I believe")
✗ Use LinkedIn buzzwords (disruptive, innovative, solution-oriented)
✗ Leave any paragraph shorter than 40 words or longer than 140 words
✗ Make up skills, certifications, or experience
✗ Fabricate company details

=== VALIDATION CHECKLIST ===
1. ✓ Does Paragraph 1 mention the specific position AND company name?
2. ✓ Does Paragraph 1 reference a specific JD requirement, tool, or responsibility?
3. ✓ Does Paragraph 2 mention 2–3 specific tools or certifications from the JD?
4. ✓ Does Paragraph 2 include an example from the resume summary?
5. ✓ Is no paragraph shorter than 40 words or longer than 140 words?
6. ✓ Does Paragraph 3 end with a specific call to action?
7. ✓ Are there NO unfilled placeholders or generic fillers?
8. ✓ Total word count: 150–220 words?
9. ✓ NO forbidden phrases (passionate, dynamic, synergy, team player)?
10. ✓ Tone: professional and direct (no apologies)?
"""


def build_cover_letter_text(
    job: Dict[str, Any],
    summary: str,
    skills: list
) -> str:
    """
    Build a high-quality fallback cover letter template when LLM is unavailable.
    This ensures a letter is still generated even if Groq API fails.
    """
    title = job.get("title", "this role")
    company = job.get("company", "your organization")
    description = job.get("description", "")[:200]
    
    # Extract first meaningful sentence or phrase from description
    desc_sample = description.split(".")[0].rstrip() if description else "the role"
    
    # Build skill references from extracted skills
    skill_list = ", ".join(skills[:3]) if skills else "technical skills"
    
    # Extract potential keywords (look for capitalized phrases, tech terms)
    potential_keywords = [word for word in description.split() if len(word) > 5 and word[0].isupper()][:3]
    keywords_str = ", ".join(potential_keywords) if potential_keywords else skill_list
    
    return f"""I am writing to express my strong interest in the {title} position at {company}. 
With my background in {skill_list} and commitment to professional excellence, I am well-positioned to contribute 
meaningfully to your team from day one.

My experience directly aligns with the core requirements of this role. I have developed expertise in {keywords_str} 
through hands-on work in {skill_list}. I understand the importance of the challenges outlined in your description—{desc_sample}—and 
I have demonstrated the ability to deliver results in similar contexts. I am particularly drawn to how this role 
emphasizes {keywords_str}, an area where I have built significant depth.

I would welcome the opportunity to discuss how my background in {skill_list} can support {company}'s goals. 
I'm confident in my ability to make an immediate impact and look forward to exploring this opportunity further."""


def generate_cover_letter(
    job: Dict[str, Any],
    summary: str,
    skills: list,
    output_dir: str,
    tailored_resume_summary: str = None
) -> str:
    """
    Generate a professional, highly tailored cover letter with structured JD parsing.
    
    Args:
        job: Job dict with title, company, description, job_url, location
        summary: Candidate's professional summary/background (fallback)
        skills: List of candidate's extracted skills from resume
        output_dir: Directory to save the cover letter .docx file
        tailored_resume_summary: Pre-tailored resume summary from Resume Tailor Agent (preferred)
    
    Returns:
        Full file path to the generated cover letter (.docx file)
    """
    try:
        # Use tailored resume summary if available; fallback to generic summary
        resume_summary = tailored_resume_summary or summary
        if not resume_summary or resume_summary.strip() == "":
            resume_summary = "strong technical background"
        
        cover_letter_text = ""
        
        # Try Groq LLM with advanced prompt first
        try:
            # Build the advanced prompt with structured JD parsing
            prompt = build_advanced_prompt(
                company_name=job.get("company", "Company"),
                job_title=job.get("title", "Position"),
                job_description=job.get("description", ""),
                tailored_resume_summary=resume_summary,
                extracted_skills=skills
            )
            
            cover_letter_text = groq_call(
                prompt=prompt,
                model="llama-3.1-8b-instant",
                max_tokens=2000,
            )
            
            # Validate output (check for minimum quality)
            if not cover_letter_text or len(cover_letter_text.split("\n")) < 2 or len(cover_letter_text) < 150:
                logger.warning("LLM output too short; using fallback template")
                cover_letter_text = build_cover_letter_text(job, resume_summary, skills)
        except RuntimeError:
            logger.warning("GROQ_API_KEY not configured; using fallback cover letter template")
            cover_letter_text = build_cover_letter_text(job, resume_summary, skills)
        except Exception as llm_error:
            logger.warning(f"Groq LLM generation failed: {llm_error}; using fallback template")
            cover_letter_text = build_cover_letter_text(job, resume_summary, skills)
        
        # Final fallback if text is still empty
        if not cover_letter_text:
            cover_letter_text = build_cover_letter_text(job, resume_summary, skills)

        # Create .docx document
        doc = Document()
        
        # Add title
        company = job.get("company", "Company")
        title = job.get("title", "Position")
        doc.add_heading(f"Cover Letter for {title} at {company}", 0)
        
        # Add cover letter content (paragraphs)
        for paragraph_text in cover_letter_text.split("\n\n"):
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text.strip())
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                # Add spacing for readability
                p.paragraph_format.space_after = Pt(6)
        
        # Sanitise filename
        sanitised_company = sanitise_filename(company)
        sanitised_title = sanitise_filename(title)
        filename = f"{sanitised_company}_{sanitised_title}_cover_letter.docx"
        
        # Ensure output directory exists
        output_path = Path(output_dir) / filename
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Save document
        doc.save(str(output_path))
        
        logger.info(f"Generated cover letter: {output_path}")
        return str(output_path)
    
    except Exception as e:
        logger.error(f"Error generating cover letter: {e}")
        return ""
